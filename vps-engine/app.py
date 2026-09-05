"""
Servidor Unificado de Procesamiento Documental, OCR e IA
DISTRIBUIDORA DE ELECTRICIDAD ENDE DEORURO S.A.
Tecnologías: FastAPI, MarkItDown, Tesseract (spa), Poppler, PyMuPDF, python-docx, pdf2docx, DeepSeek
"""

import os
import io
import re
import json
import base64
import shutil
import tempfile
import requests
from typing import Optional, List, Dict, Any
from datetime import datetime

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

import fitz  # PyMuPDF
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

import pytesseract
from PIL import Image

try:
    from markitdown import MarkItDown
    markitdown_engine = MarkItDown()
except Exception as e:
    markitdown_engine = None
    print(f"Warning MarkItDown: {e}")

try:
    from pdf2docx import Converter
except Exception as e:
    Converter = None
    print(f"Warning pdf2docx: {e}")

app = FastAPI(
    title="ENDE DEORURO - Motor Documental y de IA",
    version="2.0.0",
    description="Microservicio de extracción documental (MarkItDown/OCR), análisis DeepSeek y generación DOCX/PDF oficial."
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OUTPUT_DIR = "/tmp/ende_docs" if os.name != "nt" else os.path.join(tempfile.gettempdir(), "ende_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

OPENCODE_API_KEY = os.getenv("OPENCODE_GO_API_KEY", "sk-uiqURVX900evBUHKomZL4LjIe3L1NvILaNAcATY4oZ6rWvDMoVAt9ODP3F6Q8g97")
OPENCODE_BASE_URL = os.getenv("OPENCODE_GO_BASE_URL", "https://opencode.ai/zen/go/v1").rstrip("/")
OPENCODE_MODEL = os.getenv("OPENCODE_GO_MODEL", "deepseek-v4-flash-vision-exp")


# ============================================================================
# 1. EXTRACTORES DE TEXTO UNIVERSALES (MarkItDown, OCR Tesseract, PyMuPDF)
# ============================================================================

def extract_text_from_file_bytes(file_bytes: bytes, filename: str) -> str:
    """Extrae texto usando MarkItDown, PyMuPDF o OCR Tesseract según el tipo de archivo."""
    ext = os.path.splitext(filename)[1].lower()
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        # Intento 1: MarkItDown (Word, Excel, PowerPoint, TXT, MD, HTML)
        if markitdown_engine and ext in [".docx", ".doc", ".xlsx", ".xls", ".pptx", ".txt", ".md", ".csv", ".html"]:
            try:
                res = markitdown_engine.convert(tmp_path)
                if res and res.text_content and len(res.text_content.strip()) > 10:
                    return res.text_content.strip()
            except Exception as ex:
                print(f"MarkItDown fallback: {ex}")

        # Intento 2: PDF Digital o Escaneado
        if ext == ".pdf":
            try:
                doc = fitz.open(tmp_path)
                full_text = []
                has_text = False
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    t = page.get_text()
                    if t and len(t.strip()) > 15:
                        has_text = True
                        full_text.append(t)
                    else:
                        # Si la página no tiene texto digital, aplicar OCR Tesseract en español
                        pix = page.get_pixmap(dpi=300)
                        img = Image.open(io.BytesIO(pix.tobytes("png")))
                        ocr_t = pytesseract.image_to_string(img, lang="spa")
                        if ocr_t.strip():
                            full_text.append(ocr_t)
                
                doc.close()
                extracted = "\n\n".join(full_text).strip()
                if extracted:
                    return extracted
            except Exception as ex:
                print(f"PDF extraction error: {ex}")

        # Intento 3: Imágenes (PNG, JPG, JPEG, WEBP, TIFF)
        if ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
            try:
                img = Image.open(tmp_path)
                ocr_t = pytesseract.image_to_string(img, lang="spa")
                if ocr_t.strip():
                    return ocr_t.strip()
            except Exception as ex:
                print(f"Image OCR error: {ex}")

        # Intento 4: Decodificación UTF-8 / Latin-1
        try:
            return file_bytes.decode("utf-8")
        except:
            return file_bytes.decode("latin-1", errors="ignore")

    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except:
                pass


# ============================================================================
# 2. MOTOR DE INTELIGENCIA ARTIFICIAL (DeepSeek con Master Prompt ENDE DEORURO)
# ============================================================================

MASTER_SYSTEM_PROMPT = """# Rol y Propósito:
Eres el Asistente Técnico Oficial de Contrataciones y Adquisiciones de DISTRIBUIDORA DE ELECTRICIDAD ENDE DEORURO S.A. Tu función es transformar cualquier requerimiento, nota técnica o borrador en un documento formal de **Especificaciones Técnicas (ET) o Términos de Referencia (TdR)**, cumpliendo rigurosamente con la estructura oficial de 14 puntos de la empresa.

---

# 🧠 LÓGICA DE DETECCIÓN Y ADAPTACIÓN SEGÚN EL RUBRO (PUNTO 3)
Al procesar la solicitud, identifica la categoría para adaptar el **Punto 3 (ESPECIFICACIÓN TÉCNICA)**:

### OPCIÓN A: BIENES, HERRAMIENTAS Y EQUIPOS (Suministros)
- **Estructura:** Ficha técnica y cuadro físico/mecánico.
- **Tipo de tabla:** "BIENES_SIMPLE"
- **Formato de Tabla:**
  | No. | DESCRIPCIÓN DEL ÍTEM | CARACTERÍSTICAS / ESPECIFICACIÓN TÉCNICA | CANT. |
  | --- | -------------------- | --------------------------------------- | ----- |

### OPCIÓN B: SALUD OCUPACIONAL, MEDICINA Y SERVICIOS DE LABORATORIO
- **Estructura:** Matriz de evaluación médica y requisitos de laboratorio/consulta.
- **Tipo de tabla:** "SALUD_OCUPACIONAL"
- **Formato de Tabla:**
  | EXAMEN / SERVICIO REQUERIDO | ESPECIFICACIÓN MÍNIMA REQUERIDA | PROPUESTO / INFORMAR |
  | --------------------------- | ------------------------------- | -------------------- |

---

# 📜 ESTRUCTURA OFICIAL DEL DOCUMENTO (14 PUNTOS)
### LOS 14 PUNTOS OBLIGATORIOS:
1. **ANTECEDENTES:** Contexto operativo y normativo institucional de ENDE DEORURO S.A. (Extensión controlada: MÍNIMO 5 LÍNEAS, MÁXIMO 8 LÍNEAS).
2. **JUSTIFICACIÓN / NECESIDAD:** Justificación técnica de la necesidad operativa, continuidad del servicio eléctrico y mitigación de riesgos de accidentes basada en los ítems solicitados (Extensión controlada: MÍNIMO 5 LÍNEAS, MÁXIMO 8 LÍNEAS).
3. **ESPECIFICACIÓN TÉCNICA:** Tabla completa con todos los ítems individuales solicitados, reconociendo cantidades (incluso escritas en palabras: 'dos palas' -> Cantidad 2, 'una cinta' -> Cantidad 1), unidad ('PZA', 'ROLLO', etc.) y especificaciones técnicas detalladas con normas ASTM/IEC/ISO.
4. **CALIDAD:** Estándares normativos aplicables (ASTM/IEC/ISO), garantía técnica mínima de 12 meses y certificados del fabricante.
5. **ÁMBITO DE APLICACIÓN:** Cuadrillas técnicas, personal operativo y subestaciones de ENDE DEORURO S.A.
6. **MÉTODO DE SELECCIÓN:** Menor Precio (Art. 31 del Reglamento SBC).
7. **VIGENCIA DE LA PROPUESTA:** Validez mínima de 30 días calendario.
8. **CATEGORÍA:** Bienes / Herramientas / Salud Ocupacional.
9. **LUGAR DE ENTREGA:** Almacenes ENDE DEORURO S.A., Oruro.
10. **TIEMPO DE ENTREGA:** Máximo 30 días calendario a partir de la Orden de Compra.
11. **FORMA DE ADJUDICACIÓN:** Por ítem requerido (Art. 31 SBC).
12. **PARA LA ACEPTACIÓN DEL LOTE / SERVICIO:** Inspección y evaluación técnica de conformidad en almacén.
13. **FORMA DE PAGO:** 100% contra entrega a satisfacción, informe de conformidad y factura oficial.
14. **APLICACIÓN DE MULTAS:** 0.25% por día de retraso injustificado.

# 📌 REGLAS DE FIDELIDAD Y EXTRACCIÓN DE ÍTEMS:
1. **Extensión de Antecedentes y Justificación:** Mínimo 5 líneas y máximo 8 líneas de redacción formal para cada una.
2. **Extracción Total de Ítems:** Extrae cada bien, herramienta, material o servicio solicitado. No omitas ninguno.
3. **Reconocimiento de Cantidades:** Identifica correctamente las cantidades numéricas aun si están escritas en palabras en español (ej: "dos palas" -> Cantidad: 2, Descripción: "PALAS"; "una cinta..." -> Cantidad: 1, Descripción: "CINTA AISLANTE 1000V"; "20 alicates..." -> Cantidad: 20, Descripción: "ALICATES UNIVERSALES 8 PULGADAS").
4. **Limpieza de Descripción:** NO incluyas palabras de cantidad ("DOS", "UNA", "TRES", etc.) dentro de la descripción del ítem; trasládalas al campo "cantidad".
5. **Dimensiones y Especificaciones Técnicas:** Conserva las medidas, calibres y voltajes ("8 PULGADAS", "1000V", "6 PULGADAS") en la descripción y redacta características técnicas completas con normas ASTM/IEC/ISO para cada ítem.
6. **Copia Fiel y Coherencia:** Redacta Antecedentes y Justificación basados exclusivamente en los ítems solicitados, sin mezclar rubros diferentes.

DEBES RESPONDER EXCLUSIVAMENTE UN OBJETO JSON VÁLIDO CON ESTA ESTRUCTURA:
{
  "titulo_proceso": "ESPECIFICACIONES TÉCNICAS - ADQUISICIÓN DE ...",
  "categoria_detectada": "Bienes" | "Servicios" | "Salud Ocupacional",
  "tipo_tabla_sugerido": "BIENES_SIMPLE" | "SALUD_OCUPACIONAL" | "MATRIZ_SERVICIOS",
  "antecedentes_texto": "Texto formal de 5 a 8 líneas de antecedentes",
  "justificacion_texto": "Texto formal de 5 a 8 líneas de justificación basado en los ítems solicitados",
  "calidad_texto": "Texto de calidad",
  "ambito_aplicacion": "Texto de ámbito",
  "metodo_seleccion_texto": "Menor Precio (Art. 31 del Reglamento SBC)",
  "vigencia_propuesta_texto": "30 días calendario computables a partir de la fecha de presentación",
  "categoria_texto": "Bienes y Herramientas",
  "lugar_entrega": "Almacenes ENDE DEORURO S.A., Oruro",
  "tiempo_entrega_texto": "Máximo 30 días calendario",
  "forma_adjudicacion": "Por ítem requerido, formalizada por Orden de Compra (Art. 31 SBC)",
  "aceptacion_lote": "Inspección técnica de conformidad al momento de la entrega",
  "forma_pago_texto": "El pago se realizará contra entrega a satisfacción...",
  "multas_texto": "Multa del 0.25% por cada día de retraso",
  "items": [
    {
      "item": 1,
      "descripcion": "DESCRIPCIÓN DEL ÍTEM",
      "cantidad": 20,
      "unidad": "PZA",
      "caracteristicasTecnicas": "Especificación técnica detallada con normas aplicables"
    }
  ]
}"""

def extract_json_from_deepseek(raw: str) -> Dict[str, Any]:
    """Analizador tolerante multi-etapa con escape de strings y búsqueda directa."""
    text = raw.strip()
    if text.startswith("```json"):
        text = text[7:]
    elif text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]
    text = text.strip()
    
    # Intento 1: Parseo directo estándar
    try:
        return json.loads(text)
    except Exception:
        pass

    # Intento 2: Buscar el primer { y el último }
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        sub = text[start:end+1]
        try:
            return json.loads(sub)
        except Exception:
            pass

    # Intento 3: Reparar caracteres de escape dentro de strings
    def escape_control_chars_in_strings(s: str) -> str:
        out = []
        in_string = False
        escaped = False
        for ch in s:
            if ch == '"' and not escaped:
                in_string = not in_string
                out.append(ch)
            elif ch == '\\' and not escaped:
                escaped = True
                out.append(ch)
            elif in_string and ch == '\n':
                out.append('\\n')
            elif in_string and ch == '\r':
                pass
            elif in_string and ch == '\t':
                out.append('\\t')
            else:
                if escaped:
                    escaped = False
                out.append(ch)
        return "".join(out)

    try:
        target = text[start:end+1] if (start != -1 and end != -1) else text
        repaired = escape_control_chars_in_strings(target)
        repaired = re.sub(r",\s*(\}|\])", r"\1", repaired)
        return json.loads(repaired)
    except Exception as e:
        print(f"JSON stage 3 repair error: {e}")

    # Intento 4: Extracción por expresiones regulares
    result = {
        "titulo_proceso": "ESPECIFICACIONES TÉCNICAS - ADQUISICIÓN DE HERRAMIENTAS Y SUMINISTROS",
        "categoria_detectada": "Bienes",
        "tipo_tabla_sugerido": "BIENES_SIMPLE",
        "antecedentes_texto": "",
        "justificacion_texto": "",
        "items": []
    }
    
    m_ant = re.search(r'"antecedentes_texto"\s*:\s*"((?:[^"\\]|\\.)*)"', text)
    if m_ant:
        result["antecedentes_texto"] = m_ant.group(1).replace("\\n", "\n").replace('\\"', '"')
        
    m_just = re.search(r'"justificacion_texto"\s*:\s*"((?:[^"\\]|\\.)*)"', text)
    if m_just:
        result["justificacion_texto"] = m_just.group(1).replace("\\n", "\n").replace('\\"', '"')
        
    m_items = re.search(r'"items"\s*:\s*\[(.*?)\]\s*(?:,|\})', text, re.DOTALL)
    if m_items:
        items_str = m_items.group(1)
        raw_items = re.findall(r"\{[^{}]*\}", items_str)
        for idx, r_it in enumerate(raw_items):
            try:
                fixed_it = escape_control_chars_in_strings(r_it)
                fixed_it = re.sub(r",\s*(\}|\])", r"\1", fixed_it)
                it_obj = json.loads(fixed_it)
                result["items"].append(it_obj)
            except Exception:
                pass
                
    return result

def call_deepseek_ai(user_prompt: str) -> Dict[str, Any]:
    """Llama a DeepSeek IA y retorna el JSON estructurado de 14 puntos e ítems."""
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {OPENCODE_API_KEY}"
    }
    payload = {
        "model": OPENCODE_MODEL,
        "messages": [
            {"role": "system", "content": MASTER_SYSTEM_PROMPT},
            {"role": "user", "content": f"Requerimiento / Solicitud:\n{user_prompt}\n"}
        ],
        "temperature": 0.1,
        "max_tokens": 4096
    }
    
    res = requests.post(f"{OPENCODE_BASE_URL}/chat/completions", headers=headers, json=payload, timeout=60)
    if not res.ok:
        raise Exception(f"DeepSeek API Error [{res.status_code}]: {res.text}")
    
    raw = res.json()["choices"][0]["message"]["content"]
    return extract_json_from_deepseek(raw)


# ============================================================================
# 3. GENERADOR DE WORD (.DOCX) OFICIAL ENDE DEORURO S.A.
# ============================================================================

def create_official_tdr_docx(data: Dict[str, Any]) -> str:
    """Crea el archivo .docx oficial con los 14 puntos y la tabla de especificaciones."""
    doc = docx.Document()
    
    # Configurar márgenes de 2.5 cm
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Título Principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_title.add_run("DISTRIBUIDORA DE ELECTRICIDAD ENDE DEORURO S.A.\nESPECIFICACIONES TÉCNICAS\n\n")
    r_sub.font.bold = True
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(0, 51, 102)
    
    r_main = p_title.add_run((data.get("titulo_proceso") or "ADQUISICIÓN DE BIENES Y SUMINISTROS").upper())
    r_main.font.bold = True
    r_main.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # 14 Puntos Oficiales
    puntos = [
        (1, "ANTECEDENTES", data.get("antecedentes_texto")),
        (2, "JUSTIFICACIÓN / NECESIDAD", data.get("justificacion_texto")),
        (3, "ESPECIFICACIÓN TÉCNICA", None), # Tabla
        (4, "CALIDAD", data.get("calidad_texto")),
        (5, "ÁMBITO DE APLICACIÓN", data.get("ambito_aplicacion")),
        (6, "MÉTODO DE SELECCIÓN", data.get("metodo_seleccion_texto")),
        (7, "VIGENCIA DE LA PROPUESTA", data.get("vigencia_propuesta_texto")),
        (8, "CATEGORÍA", data.get("categoria_texto")),
        (9, "LUGAR DE ENTREGA", data.get("lugar_entrega")),
        (10, "TIEMPO DE ENTREGA", data.get("tiempo_entrega_texto")),
        (11, "FORMA DE ADJUDICACIÓN", data.get("forma_adjudicacion")),
        (12, "PARA LA ACEPTACIÓN DEL LOTE / SERVICIO", data.get("aceptacion_lote")),
        (13, "FORMA DE PAGO", data.get("forma_pago_texto")),
        (14, "APLICACIÓN DE MULTAS", data.get("multas_texto")),
    ]
    
    for num, titulo, contenido in puntos:
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(8)
        p_head.paragraph_format.space_after = Pt(4)
        r_head = p_head.add_run(f"{num}. {titulo}")
        r_head.font.bold = True
        r_head.font.size = Pt(11)
        
        if num == 3:
            # Tabla de Ítems
            items = data.get("items", [])
            table = doc.add_table(rows=1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "No."
            hdr_cells[1].text = "DESCRIPCIÓN DEL ÍTEM"
            hdr_cells[2].text = "CARACTERÍSTICAS / ESPECIFICACIÓN TÉCNICA"
            hdr_cells[3].text = "CANT."
            
            for cell in hdr_cells:
                shading_elm = parse_xml(r'<w:shd {} w:fill="ECEEF0"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(10)
                        
            for it in items:
                row_cells = table.add_row().cells
                row_cells[0].text = str(it.get("item", ""))
                row_cells[1].text = str(it.get("descripcion", "")).upper()
                row_cells[2].text = str(it.get("caracteristicasTecnicas") or it.get("especificacionMinima") or "")
                row_cells[3].text = f"{it.get('cantidad', 1)} {it.get('unidad', 'PZA')}"
                
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for cell in row_cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9.5)
                            
            doc.add_paragraph()
        else:
            if contenido:
                p_text = doc.add_paragraph()
                p_text.paragraph_format.space_after = Pt(6)
                p_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                r_text = p_text.add_run(contenido)
                r_text.font.size = Pt(10)
                
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_name = f"TDR_ENDE_DEORURO_{timestamp}.docx"
    file_path = os.path.join(OUTPUT_DIR, file_name)
    doc.save(file_path)
    return file_name


# ============================================================================
# 4. ENDPOINTS API (FastAPI)
# ============================================================================

class ProcessRequest(BaseModel):
    insumoTexto: Optional[str] = None
    documentText: Optional[str] = None
    imageBase64: Optional[str] = None
    nombreArchivo: Optional[str] = None
    adquisicion: Optional[Dict[str, Any]] = None

@app.post("/api/procesar-documento")
async def procesar_documento_json(req: ProcessRequest):
    """Procesa requerimientos enviados en formato JSON (texto, base64 o metadatos)."""
    extracted_text = ""
    
    if req.imageBase64:
        b64_data = req.imageBase64.split(",")[-1]
        file_bytes = base64.b64decode(b64_data)
        fname = req.nombreArchivo or "archivo.pdf"
        extracted_text = extract_text_from_file_bytes(file_bytes, fname)
    elif req.documentText:
        extracted_text = req.documentText
    elif req.insumoTexto:
        extracted_text = req.insumoTexto
        
    if not extracted_text.strip():
        extracted_text = "Adquisición de bienes y herramientas operativas para ENDE DEORURO S.A."

    # 1. Analizar con DeepSeek IA
    ai_data = call_deepseek_ai(extracted_text)
    
    # 2. Generar Word .docx Oficial
    docx_file = create_official_tdr_docx(ai_data)
    
    return {
        "success": True,
        "data": ai_data,
        "extracted_text_preview": extracted_text[:500],
        "docx_file": docx_file,
        "download_docx": f"/download/{docx_file}",
        "download_pdf": f"/download/{docx_file}"
    }

@app.post("/api/procesar-archivo")
async def procesar_archivo_upload(
    file: UploadFile = File(...),
    texto_adicional: Optional[str] = Form(None)
):
    """Procesa archivos subidos directamente por multipart/form-data."""
    contents = await file.read()
    filename = file.filename or "archivo"
    extracted_text = extract_text_from_file_bytes(contents, filename)
    
    if texto_adicional:
        extracted_text = f"{texto_adicional}\n\n{extracted_text}"
        
    if not extracted_text.strip():
        extracted_text = "Adquisición de bienes y herramientas operativas para ENDE DEORURO S.A."

    ai_data = call_deepseek_ai(extracted_text)
    docx_file = create_official_tdr_docx(ai_data)
    
    return {
        "success": True,
        "data": ai_data,
        "extracted_text_preview": extracted_text[:500],
        "docx_file": docx_file,
        "download_docx": f"/download/{docx_file}",
        "download_pdf": f"/download/{docx_file}"
    }

@app.post("/api/generar-especificaciones")
async def generar_especificaciones(payload: Dict[str, Any]):
    """Compatibilidad con el endpoint anterior de especificaciones."""
    docx_file = create_official_tdr_docx(payload)
    return {
        "status": "success",
        "message": "Especificaciones técnicas generadas exitosamente para ENDE DEORURO S.A.",
        "docx_file": docx_file,
        "download_docx": f"/download/{docx_file}",
        "download_pdf": f"/download/{docx_file}"
    }

@app.post("/api/convertir-plantilla")
async def convertir_plantilla(
    file: UploadFile = File(...),
    fk_carpeta: int = Form(1),
    nombre_plantilla: Optional[str] = Form(None)
):
    """
    Convierte una plantilla subida por el usuario (PDF o DOCX) en un molde editable con pdf2docx y MarkItDown.
    """
    filename = file.filename or "plantilla.pdf"
    ext = os.path.splitext(filename)[1].lower()
    contents = await file.read()
    
    extracted_text = extract_text_from_file_bytes(contents, filename)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    if ext == ".pdf" and Converter:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
            tmp_pdf.write(contents)
            pdf_path = tmp_pdf.name
            
        out_docx_name = f"plantilla_carpeta_{fk_carpeta}_{timestamp}.docx"
        out_docx_path = os.path.join(OUTPUT_DIR, out_docx_name)
        try:
            cv = Converter(pdf_path)
            cv.convert(out_docx_path)
            cv.close()
        except Exception as e:
            print(f"Error en pdf2docx convert: {e}")
            out_docx_name = filename
        finally:
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
    else:
        out_docx_name = f"plantilla_carpeta_{fk_carpeta}_{timestamp}{ext}"
        out_docx_path = os.path.join(OUTPUT_DIR, out_docx_name)
        with open(out_docx_path, "wb") as f:
            f.write(contents)

    # Análisis inteligente del documento con DeepSeek IA para autocompletar propósito, dependencias y pasos
    ai_analysis = {}
    if extracted_text.strip():
        prompt = f"""Analiza esta plantilla oficial subida para ENDE DEORURO S.A. ({filename}) y responde en JSON estricto:
{{
  "titulo_sugerido": "Título Formal Institucional Corto",
  "descripcion": "Descripción clara de 1 línea para usuarios no técnicos",
  "que_hace": "Explicación sencilla y clara de para qué sirve este documento y qué trámite realiza",
  "de_quien_depende": "Explicación clara de qué carpetas, documentos o datos previos necesita",
  "pasos": [
    "Paso 1: Primer paso recomendado",
    "Paso 2: Segundo paso recomendado",
    "Paso 3: Tercer paso recomendado"
  ]
}}

Texto extraído de la plantilla:
{extracted_text[:4000]}"""
        try:
            ai_analysis = call_deepseek_ai(prompt)
        except Exception as e:
            print(f"Error analizando plantilla con IA: {e}")

    return {
        "success": True,
        "fk_carpeta": fk_carpeta,
        "nombre_archivo": out_docx_name,
        "download_docx": f"/download/{out_docx_name}",
        "analisis_ia": ai_analysis,
        "texto_extraido_preview": extracted_text[:1000],
        "total_caracteres": len(extracted_text)
    }

@app.post("/api/procesar-proforma-ocr")
async def procesar_proforma_ocr(
    file: Optional[UploadFile] = File(None),
    texto: Optional[str] = Form(None),
    imageBase64: Optional[str] = Form(None)
):
    """
    Extrae datos de cotizaciones / proformas escaneadas (Carpetas 2, 3 y 4) usando Tesseract OCR + MarkItDown + DeepSeek.
    """
    extracted_text = ""
    if file:
        contents = await file.read()
        extracted_text = extract_text_from_file_bytes(contents, file.filename or "proforma.pdf")
    elif imageBase64:
        b64_data = imageBase64.split(",")[-1]
        file_bytes = base64.b64decode(b64_data)
        extracted_text = extract_text_from_file_bytes(file_bytes, "proforma.png")
    elif texto:
        extracted_text = texto

    prompt = f"""Analiza esta cotización/proforma para ENDE DEORURO S.A. y extrae los datos clave en JSON estricto:
{{
  "proveedor": "Nombre o Razón Social del Proveedor",
  "nit": "NIT si figura",
  "monto_total_bs": 12500.0,
  "plazo_entrega_dias": 30,
  "validez_propuesta_dias": 30,
  "cumple_especificaciones": true,
  "items": [
    {{
      "item": 1,
      "descripcion": "DESCRIPCIÓN DEL ÍTEM",
      "cantidad": 10,
      "precio_unitario_bs": 100.0,
      "precio_total_bs": 1000.0
    }}
  ]
}}

Texto de la proforma:
{extracted_text}"""

    res = call_deepseek_ai(prompt)
    return {
        "success": True,
        "data": res,
        "extracted_text_preview": extracted_text[:500]
    }

@app.post("/api/procesar-solicitud-inicio")
async def procesar_solicitud_inicio(req: ProcessRequest):
    """Genera la Solicitud de Inicio (Carpeta 5) con DeepSeek IA y motor VPS."""
    extracted = req.documentText or req.insumoTexto or ""
    adq = req.adquisicion or {}
    
    prompt = f"""Genera la SOLICITUD DE INICIO DE CONTRATACIÓN formal para ENDE DEORURO S.A. en JSON:
{{
  "objeto": "SOLICITUD DE INICIO DEL PROCESO DE CONTRATACIÓN PARA {adq.get('titulo_proceso', '')}",
  "parrafo1": "Por medio de la presente, me dirijo a su autoridad para solicitar formalmente el inicio del proceso...",
  "justificacion": "La presente solicitud responde a la necesidad técnica de...",
  "partida": "{adq.get('partida_presupuestaria', '39500 - Herramientas Menores')}",
  "prevision_presupuesto": {adq.get('prevision_presupuesto', 0)},
  "fecha_solicitud": "{datetime.now().strftime('%d de %B de %Y')}"
}}

Contexto:
{extracted}"""

    res = call_deepseek_ai(prompt)
    return {"success": True, "data": res}

@app.post("/api/procesar-form-s2")
async def procesar_form_s2(req: ProcessRequest):
    """Genera el Formulario S-2 (Carpeta 6) con DeepSeek IA."""
    adq = req.adquisicion or {}
    prompt = f"""Genera el FORMULARIO S-2 (Solicitud de Cotización / Propuesta) para ENDE DEORURO S.A. en JSON:
{{
  "senores": "PROPONENTE / PROVEEDOR",
  "tiempo_entrega": "30 días calendario",
  "validez_oferta": "30 días calendario",
  "observaciones": "SE ADJUNTA ESPECIFICACIONES TÉCNICAS OFICIALES",
  "nota_adicional": "ADJUNTAR FOTOCOPIA SIMPLE DE SU RNC - NIT Y REGISTRO EN FUNDEMPRESA/SEPREC"
}}

Proceso: {adq.get('titulo_proceso', '')}"""

    res = call_deepseek_ai(prompt)
    return {"success": True, "data": res}

@app.post("/api/procesar-informe-conformidad")
async def procesar_informe_conformidad(req: ProcessRequest):
    """Genera el Informe Técnico de Conformidad (Carpeta 7) con DeepSeek IA."""
    adq = req.adquisicion or {}
    prompt = f"""Genera el INFORME TÉCNICO DE CONFORMIDAD Y RECEPCIÓN DEFINITIVA para ENDE DEORURO S.A. en JSON:
{{
  "conclusion": "Habiéndose verificado el cumplimiento estricto del 100% de las especificaciones técnicas requeridas, cantidades y calidad de los bienes/servicios entregados, se emite la CONFORMIDAD TÉCNICA Y RECEPCIÓN DEFINITIVA sin observaciones.",
  "recomendacion": "Se recomienda proceder al pago correspondiente de acuerdo al contrato/orden de compra y reglamentación interna de la empresa."
}}

Proceso: {adq.get('titulo_proceso', '')}"""

    res = call_deepseek_ai(prompt)
    return {"success": True, "data": res}

@app.post("/api/procesar-memo-pago")
async def procesar_memo_pago(req: ProcessRequest):
    """Genera el Memorándum de Pago (Carpeta 8) con DeepSeek IA."""
    adq = req.adquisicion or {}
    prompt = f"""Genera el MEMORÁNDUM DE SOLICITUD DE DESEMBOLSO Y PAGO para ENDE DEORURO S.A. en JSON:
{{
  "objeto": "SOLICITUD DE PAGO CORRESPONDIENTE AL PROCESO {adq.get('titulo_proceso', '')}",
  "monto_literal": "{adq.get('prevision_presupuesto', 0)} Bolivianos",
  "documentos_adjuntos": [
    "Nota de Entrega / Acta de Recepción Definitiva",
    "Informe Técnico de Conformidad Aprobado",
    "Factura Comercial Original",
    "Orden de Compra / Contrato Administrativo",
    "Formulario S-2 de Adjudicación"
  ]
}}"""

class AssistantFolderRequest(BaseModel):
    peticion_usuario: str
    adquisicion: Optional[Dict[str, Any]] = None
    carpetas_existentes: Optional[List[Dict[str, Any]]] = None

@app.post("/api/asistente-carpeta")
async def asistente_carpeta(req: AssistantFolderRequest):
    """
    Asistente experto de ENDE DEORURO que traduce peticiones no técnicas en la estructura óptima de carpeta.
    """
    adq = req.adquisicion or {}
    pet = req.peticion_usuario or "Crear nueva carpeta para el proceso"
    
    prompt = f"""Eres el Asistente Experto en Contrataciones de ENDE DEORURO S.A.
El usuario no es técnico y te pide lo siguiente con sus propias palabras:
"{pet}"

Para el proceso: "{adq.get('titulo_proceso', 'Adquisición Institucional')}".
Analiza la solicitud y devuelve en JSON estricto:
{{
  "nombre_carpeta": "Título Formal e Institucional Corto (ej. Acta de Apertura de Sobres)",
  "descripcion_clara": "Explicación breve de 1 línea de qué es este documento",
  "que_hace": "Explicación en lenguaje sencillo y claro de lo que se tramita o redacta en esta fase",
  "de_quien_depende": "Explicación clara de qué carpetas o documentos previos se necesita tener listos",
  "pasos": [
    "Paso 1: Explicación sencilla del primer paso",
    "Paso 2: Explicación sencilla del segundo paso",
    "Paso 3: Explicación sencilla del tercer paso"
  ],
  "tipo_generacion": "IA",
  "borrador_contenido": "Texto inicial redactado formalmente para ENDE DEORURO S.A. listo para ser usado..."
}}"""


# ============================================================================
# 5. MOTOR INTELIGENTE DE PLANTILLAS DOCX (Smart DOCX Filler sin etiquetas)
# ============================================================================

def inspect_docx_file(docx_path: str) -> dict:
    doc = docx.Document(docx_path)
    paragraphs_info = []
    blank_pattern = re.compile(r'(__{2,}|\[[\w\s\.\,\-]+\]|\.{3,})')

    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        matches = blank_pattern.findall(text)
        has_colon_prompt = ":" in text and (len(text.split(":", 1)[1].strip()) == 0 or bool(matches))
        paragraphs_info.append({
            "index": idx,
            "text": text,
            "style": p.style.name if p.style else None,
            "placeholders": matches,
            "is_fillable": bool(matches) or has_colon_prompt
        })

    tables_info = []
    for t_idx, table in enumerate(doc.tables):
        rows_data = []
        for r in table.rows:
            row_cells = [cell.text.strip().replace('\n', ' ') for cell in r.cells]
            rows_data.append(row_cells)
        headers = rows_data[0] if rows_data else []
        tables_info.append({
            "table_index": t_idx,
            "total_rows": len(table.rows),
            "total_columns": len(table.columns),
            "headers": headers,
            "sample_first_row": rows_data[1] if len(rows_data) > 1 else []
        })

    return {
        "file": os.path.basename(docx_path),
        "total_paragraphs": len(doc.paragraphs),
        "fillable_paragraphs_count": sum(1 for p in paragraphs_info if p["is_fillable"]),
        "paragraphs": paragraphs_info,
        "tables": tables_info
    }

def fill_smart_docx_file(template_path: str, output_path: str, data: dict) -> dict:
    doc = docx.Document(template_path)
    stats = {"replaced_placeholders": 0, "updated_sections": 0, "updated_tables": 0}

    def preserve_run_formatting(target_run, source_run):
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        if source_run.font.color and source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline

    def replace_text_in_paragraph(paragraph, old_text, new_text):
        full_text = paragraph.text
        if old_text not in full_text:
            return False
        template_run = paragraph.runs[0] if paragraph.runs else None
        new_full_text = full_text.replace(old_text, new_text)
        p_element = paragraph._p
        for r in paragraph.runs:
            p_element.remove(r._r)
        new_run = paragraph.add_run(new_full_text)
        if template_run:
            preserve_run_formatting(new_run, template_run)
        return True

    def set_paragraph_content(paragraph, new_text):
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return
        template_run = paragraph.runs[0]
        p_element = paragraph._p
        for r in paragraph.runs:
            p_element.remove(r._r)
        new_run = paragraph.add_run(new_text)
        preserve_run_formatting(new_run, template_run)

    replacements = data.get("replacements", {})
    if replacements:
        for p in doc.paragraphs:
            for old_val, new_val in replacements.items():
                if old_val in p.text:
                    if replace_text_in_paragraph(p, old_val, str(new_val)):
                        stats["replaced_placeholders"] += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for old_val, new_val in replacements.items():
                            if old_val in p.text:
                                if replace_text_in_paragraph(p, old_val, str(new_val)):
                                    stats["replaced_placeholders"] += 1

    sections_data = data.get("sections", {})
    if sections_data:
        total_p = len(doc.paragraphs)
        for i, p in enumerate(doc.paragraphs):
            p_text = p.text.strip().upper()
            for sec_name, new_content in sections_data.items():
                sec_upper = sec_name.strip().upper()
                if sec_upper in p_text and (len(p_text) - len(sec_upper) < 15):
                    for j in range(i + 1, min(i + 4, total_p)):
                        next_p = doc.paragraphs[j]
                        if next_p.text.strip():
                            set_paragraph_content(next_p, str(new_content))
                            stats["updated_sections"] += 1
                            break

    tables_data = data.get("tables", [])
    for t_spec in tables_data:
        t_idx = t_spec.get("table_index", 0)
        if t_idx < len(doc.tables):
            table = doc.tables[t_idx]
            new_rows = t_spec.get("rows", [])
            replace_mode = t_spec.get("mode", "append")
            if replace_mode == "direct_cells":
                for r_idx, c_idx, val in t_spec.get("cells", []):
                    if r_idx < len(table.rows) and c_idx < len(table.rows[r_idx].cells):
                        table.rows[r_idx].cells[c_idx].text = str(val)
                stats["updated_tables"] += 1
            elif replace_mode == "append":
                for row_data in new_rows:
                    row = table.add_row()
                    for c_idx, val in enumerate(row_data):
                        if c_idx < len(row.cells):
                            row.cells[c_idx].text = str(val)
                stats["updated_tables"] += 1

    doc.save(output_path)
    return stats

@app.post("/api/docx/inspect")
async def api_inspect_docx(file: UploadFile = File(...)):
    """Inspecciona cualquier plantilla DOCX y devuelve su estructura editable."""
    contents = await file.read()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(contents)
        tmp_path = tmp.name
    try:
        structure = inspect_docx_file(tmp_path)
        return {"success": True, "data": structure}
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

@app.post("/api/docx/smart-fill")
async def api_smart_fill_docx(file: UploadFile = File(...), data_json: str = Form(...)):
    """Rellena cualquier plantilla DOCX usando datos JSON sin requerir tags predefinidos."""
    try:
        data = json.loads(data_json)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"JSON inválido: {e}")

    contents = await file.read()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(contents)
        tmp_path = tmp.name

    out_name = f"filled_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    out_path = os.path.join(OUTPUT_DIR, out_name)

    try:
        stats = fill_smart_docx_file(tmp_path, out_path, data)
        return {
            "success": True,
            "filename": out_name,
            "download_url": f"/download/{out_name}",
            "stats": stats
        }
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

@app.post("/api/pdf-to-docx")
async def pdf_to_docx_convert(file: UploadFile = File(...)):
    """Convierte un PDF a DOCX editable con pdf2docx."""
    if not Converter:
        raise HTTPException(status_code=500, detail="pdf2docx no está disponible en este servidor.")
        
    contents = await file.read()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        tmp_pdf.write(contents)
        pdf_path = tmp_pdf.name
        
    out_docx_name = f"converted_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    out_docx_path = os.path.join(OUTPUT_DIR, out_docx_name)
    
    try:
        cv = Converter(pdf_path)
        cv.convert(out_docx_path)
        cv.close()
        return {
            "success": True,
            "docx_file": out_docx_name,
            "download_docx": f"/download/{out_docx_name}"
        }
    finally:
        if os.path.exists(pdf_path):
            os.remove(pdf_path)

@app.get("/download/{filename}")
async def download_file(filename: str):
    """Descarga de documentos generados."""
    file_path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Archivo no encontrado.")
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

@app.get("/health")
def health_check():
    return {
        "status": "online",
        "service": "ENDE DEORURO Document & AI Engine",
        "ocr_languages": pytesseract.get_languages() if hasattr(pytesseract, "get_languages") else ["spa", "eng"],
        "markitdown": markitdown_engine is not None,
        "pdf2docx": Converter is not None
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8080)
